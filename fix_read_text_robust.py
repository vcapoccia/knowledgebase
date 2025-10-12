"""
fix_read_text_robust.py - Patch per _read_text con gestione errori avanzata

Applica questa patch a worker/worker_tasks.py per:
1. Gestire errori LibreOffice
2. Supportare più formati
3. Pulire byte NULL da PostgreSQL
4. Gestire file corrotti gracefully
5. Logging dettagliato errori

Uso:
    python3 fix_read_text_robust.py
"""

ROBUST_READ_TEXT = '''def _read_text(path: str) -> str:
    """
    Estrae testo da vari formati con gestione errori robusta.
    
    Supporto:
    - Testo: .txt, .md, .csv, .log
    - PDF: .pdf (via pdftotext)
    - Word: .docx (via python-docx), .doc (via LibreOffice)
    - Excel: .xls, .xlsx (via LibreOffice)
    - PowerPoint: .ppt, .pptx (via LibreOffice)
    - Altri: .odt, .ods, .odp, .rtf
    
    Returns:
        str: Testo estratto (vuoto se impossibile)
    
    Raises:
        RuntimeError: Solo per errori critici
    """
    import logging
    log = logging.getLogger("worker")
    
    ext = os.path.splitext(path)[1].lower()
    
    # Skip formati non supportati
    UNSUPPORTED_EXTENSIONS = {
        '.mpp', '.vsd', '.mdb', '.accdb',  # Microsoft Project, Visio, Access
        '.zip', '.rar', '.7z', '.tar', '.gz',  # Archivi
        '.exe', '.dll', '.so',  # Binari
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg',  # Immagini
        '.mp3', '.mp4', '.avi', '.mov', '.wav',  # Media
    }
    
    if ext in UNSUPPORTED_EXTENSIONS:
        log.debug(f"Formato non supportato (skip): {ext} - {os.path.basename(path)}")
        return ""  # Ritorna vuoto invece di errore
    
    try:
        # ===== TESTO PURO =====
        if ext in (".txt", ".md", ".csv", ".log", ".ini", ".conf", ".xml", ".json"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return _clean_text(content)
        
        # ===== PDF =====
        if ext == ".pdf":
            return _pdftotext_safe(path)
        
        # ===== DOCX (priorità python-docx) =====
        if ext == ".docx":
            return _extract_docx(path)
        
        # ===== ALTRI FORMATI OFFICE =====
        if ext in (".doc", ".xls", ".xlsx", ".ppt", ".pptx", ".odt", ".ods", ".odp", ".rtf"):
            return _libreoffice_convert_safe(path)
        
        # Formato sconosciuto ma proviamo come testo
        log.warning(f"Formato sconosciuto, provo come testo: {ext}")
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return _clean_text(content)
        except Exception:
            return ""
    
    except Exception as e:
        log.error(f"Errore estrazione {os.path.basename(path)}: {e}")
        return ""  # Ritorna vuoto invece di propagare errore


def _clean_text(text: str) -> str:
    """
    Pulisce testo da caratteri problematici per PostgreSQL.
    
    Rimuove:
    - NULL bytes (0x00) che causano errore PostgreSQL
    - Caratteri di controllo
    - Spazi multipli
    """
    if not text:
        return ""
    
    # Rimuovi NULL bytes (CRITICO per PostgreSQL)
    text = text.replace('\\x00', '')
    text = text.replace('\\0', '')
    
    # Rimuovi altri caratteri di controllo problematici (opzionale)
    import re
    text = re.sub(r'[\\x00-\\x08\\x0B\\x0C\\x0E-\\x1F\\x7F]', '', text)
    
    # Normalizza spazi multipli
    text = re.sub(r'\\s+', ' ', text)
    
    # Trim
    text = text.strip()
    
    return text


def _pdftotext_safe(path: str) -> str:
    """Estrazione PDF con gestione errori"""
    import logging
    log = logging.getLogger("worker")
    
    try:
        out = subprocess.run(
            ["pdftotext", "-layout", "-nopgbrk", path, "-"],
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE, 
            text=True,
            timeout=30  # Timeout 30s per PDF grandi
        )
        
        if out.returncode != 0:
            log.warning(f"pdftotext fallito per {os.path.basename(path)}: {out.stderr}")
            return ""
        
        return _clean_text(out.stdout)
    
    except subprocess.TimeoutExpired:
        log.error(f"pdftotext timeout su {os.path.basename(path)}")
        return ""
    except Exception as e:
        log.error(f"pdftotext errore su {os.path.basename(path)}: {e}")
        return ""


def _extract_docx(path: str) -> str:
    """
    Estrazione DOCX con fallback a LibreOffice.
    """
    import logging
    log = logging.getLogger("worker")
    
    try:
        from docx import Document
        
        doc = Document(path)
        
        # Estrai paragrafi
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Estrai tabelle
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        # Combina
        all_text = "\\n".join(paragraphs + tables_text)
        
        return _clean_text(all_text)
    
    except ImportError:
        # python-docx non disponibile, usa LibreOffice
        log.debug(f"python-docx non disponibile, uso LibreOffice per {os.path.basename(path)}")
        return _libreoffice_convert_safe(path)
    
    except Exception as e:
        # DOCX corrotto o problematico, prova LibreOffice
        log.warning(f"python-docx fallito su {os.path.basename(path)}: {e}, provo LibreOffice")
        return _libreoffice_convert_safe(path)


def _libreoffice_convert_safe(path: str) -> str:
    """
    Conversione LibreOffice con gestione errori robusta.
    """
    import tempfile
    import logging
    log = logging.getLogger("worker")
    
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Converti in TXT
            result = subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "txt:Text",
                "--outdir", tmpdir,
                path
            ], 
            capture_output=True, 
            text=True, 
            timeout=60)
            
            # Se conversione fallisce, ritorna vuoto invece di errore
            if result.returncode != 0:
                log.warning(f"LibreOffice fallito su {os.path.basename(path)}: {result.stderr}")
                return ""
            
            # Cerca file TXT generato
            basename = os.path.splitext(os.path.basename(path))[0]
            
            # LibreOffice può generare nomi diversi, cerca tutti i .txt
            txt_files = [f for f in os.listdir(tmpdir) if f.endswith('.txt')]
            
            if not txt_files:
                log.warning(f"LibreOffice non ha generato TXT per {os.path.basename(path)}")
                return ""
            
            # Prendi il primo (o quello che matcha il basename)
            txt_file = None
            for f in txt_files:
                if basename.lower() in f.lower():
                    txt_file = f
                    break
            
            if not txt_file:
                txt_file = txt_files[0]
            
            txt_path = os.path.join(tmpdir, txt_file)
            
            # Leggi contenuto
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            
            return _clean_text(content)
    
    except subprocess.TimeoutExpired:
        log.error(f"LibreOffice timeout su {os.path.basename(path)}")
        return ""
    
    except Exception as e:
        log.error(f"LibreOffice errore su {os.path.basename(path)}: {e}")
        return ""
'''


def apply_patch(worker_file="worker/worker_tasks.py"):
    """Applica patch al file worker"""
    import os
    import re
    from datetime import datetime
    
    if not os.path.exists(worker_file):
        print(f"❌ File non trovato: {worker_file}")
        return False
    
    # Backup
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_file = f"{worker_file}.backup.{timestamp}"
    
    with open(worker_file, "r") as f:
        content = f.read()
    
    with open(backup_file, "w") as f:
        f.write(content)
    
    print(f"💾 Backup: {backup_file}")
    
    # Trova vecchia _read_text
    pattern = r'def _read_text\(path: str\) -> str:.*?(?=\ndef [a-z_]|\Z)'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        print("❌ Funzione _read_text non trovata")
        return False
    
    old_function = match.group(0)
    print(f"✅ Trovata _read_text ({len(old_function)} char)")
    
    # Sostituisci
    new_content = content.replace(old_function, ROBUST_READ_TEXT)
    
    # Salva
    output_file = f"{worker_file}.ROBUST"
    with open(output_file, "w") as f:
        f.write(new_content)
    
    print(f"✅ File aggiornato: {output_file}")
    print()
    print("🔍 Verifica:")
    print(f"   diff {worker_file} {output_file} | head -50")
    print()
    print("✅ Applica:")
    print(f"   mv {output_file} {worker_file}")
    print()
    print("🔄 Rebuild:")
    print("   docker compose build worker")
    print("   docker compose restart worker")
    
    return True


if __name__ == "__main__":
    print("🔧 PATCH: _read_text Robusta con Gestione Errori")
    print("=" * 60)
    print()
    
    apply_patch()
