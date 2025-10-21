# worker/worker_tasks.py - MULTI-MODEL + GPU ACCELERATION
import os
import json
import subprocess
import logging
import uuid
import hashlib
from typing import Dict, Any, List, Optional
from datetime import datetime

import redis
from redis import Redis

import psycopg
from psycopg.rows import dict_row

import meilisearch
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

# ===== GPU DETECTION =====
import torch

DEVICE = None
GPU_AVAILABLE = False

try:
    if torch.cuda.is_available():
        DEVICE = 'cuda'
        GPU_AVAILABLE = True
        logging.info(f"🎮 Worker GPU DETECTED: {torch.cuda.get_device_name(0)}")
        logging.info(f"🎮 CUDA Version: {torch.version.cuda}")
    else:
        DEVICE = 'cpu'
        logging.info("💻 GPU non disponibile, worker usa CPU")
except Exception as e:
    DEVICE = 'cpu'
    logging.warning(f"⚠️ Errore detection GPU: {e}, worker usa CPU")

# Batch size ottimizzato per device
BATCH_SIZE = 48 if GPU_AVAILABLE else 16
logging.info(f"📦 Worker batch size: {BATCH_SIZE} ({'GPU' if GPU_AVAILABLE else 'CPU'} optimized)")

# ===== ENV =====
REDIS_URL = os.getenv("REDIS_URL", "redis://redis:6379/0")
RQ_QUEUE = os.getenv("RQ_QUEUE", "kb_ingestion")

POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres")
POSTGRES_DB = os.getenv("POSTGRES_DB", "kb")
POSTGRES_USER = os.getenv("POSTGRES_USER", "kbuser")
POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD", "kbpass")

MEILI_URL = os.getenv("MEILI_URL", "http://meili:7700")
MEILI_MASTER_KEY = os.getenv("MEILI_MASTER_KEY", os.getenv("MEILI_KEY", "change_me_meili_key"))

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

KB_ROOT = os.getenv("KB_ROOT", "/mnt/kb")
MEILI_INDEX = "kb_docs"

Q_REDIS_KEY_PROGRESS = "kb:progress"
Q_REDIS_KEY_FAILED = "kb:failed_docs"
Q_REDIS_KEY_CURRENT_DOC = "kb:current_doc"
Q_REDIS_KEY_PROCESSING_LOG = "kb:processing_log"
Q_REDIS_KEY_STATS = "kb:stats"

# Configurazione modelli
MODEL_CONFIGS = {
    "sentence-transformer": {
        "name": "all-MiniLM-L6-v2",
        "dimension": 384,
        "collection_prefix": "kb_st",
        "type": "transformers"
    },
    "llama3": {
        "name": "llama3",
        "dimension": 4096,
        "collection_prefix": "kb_llama3",
        "type": "ollama"
    },
    "mistral": {
        "name": "mistral",
        "dimension": 4096,
        "collection_prefix": "kb_mistral",
        "type": "ollama"
    }
}

log = logging.getLogger("worker")

# ===== GPU Memory Management =====
def clear_gpu_cache():
    """Libera memoria GPU cache"""
    if GPU_AVAILABLE:
        try:
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        except Exception as e:
            log.warning(f"Errore clear GPU cache: {e}")

# ===== Helpers =====
def rconn() -> Redis:
    return redis.from_url(REDIS_URL, decode_responses=True)

def meili_client() -> meilisearch.Client:
    return meilisearch.Client(MEILI_URL, MEILI_MASTER_KEY)

def qdrant_client() -> QdrantClient:
    return QdrantClient(url=QDRANT_URL)

def pg_conn():
    dsn = f"host={POSTGRES_HOST} dbname={POSTGRES_DB} user={POSTGRES_USER} password={POSTGRES_PASSWORD}"
    return psycopg.connect(dsn, autocommit=True, row_factory=dict_row)

def ensure_pg_schema():
    """Crea schema PostgreSQL con metadati avanzati"""
    with pg_conn() as conn, conn.cursor() as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY,
            path TEXT NOT NULL,
            title TEXT,
            content TEXT,
            mtime TIMESTAMP DEFAULT NOW(),
            
            -- Metadati base
            ext TEXT,
            
            -- Metadati strutturati (da metadata_extractor)
            area TEXT,
            anno TEXT,
            cliente TEXT,
            oggetto TEXT,
            tipo_doc TEXT,
            codice_appalto TEXT,
            categoria TEXT,
            descrizione_oggetto TEXT,
            versione TEXT
        );
        """)
        
        # Indici per performance
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_area ON documents(area);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_anno ON documents(anno);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_cliente ON documents(cliente);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_oggetto ON documents(oggetto);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_categoria ON documents(categoria);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_ext ON documents(ext);")

def _generate_point_id(doc_id: str, chunk_idx: int) -> str:
    """
    Genera UUID deterministico per point ID Qdrant.
    Usa hash MD5 per garantire stesso ID per stesso documento+chunk.
    """
    unique_str = f"{doc_id}_chunk_{chunk_idx}"
    namespace = uuid.UUID('6ba7b810-9dad-11d1-80b4-00c04fd430c8')
    point_uuid = uuid.uuid5(namespace, unique_str)
    return str(point_uuid)

def _push_failed(rc: Redis, item: Dict[str, Any]):
    """Aggiungi documento fallito alla lista"""
    rc.lpush(Q_REDIS_KEY_FAILED, json.dumps({**item, "ts": datetime.now().timestamp()}))

def _set_progress(rc: Redis, running: bool, done: int, total: int, stage: str):
    """Aggiorna progress generale"""
    rc.set(Q_REDIS_KEY_PROGRESS, json.dumps({
        "running": running, "done": done, "total": total, "stage": stage
    }))

def _set_current_doc(rc: Redis, doc_info: Optional[Dict[str, Any]]):
    """Imposta documento corrente in elaborazione"""
    if doc_info:
        rc.set(Q_REDIS_KEY_CURRENT_DOC, json.dumps({**doc_info, "ts": datetime.now().timestamp()}))
    else:
        rc.delete(Q_REDIS_KEY_CURRENT_DOC)

def _add_processing_log(rc: Redis, entry: Dict[str, Any]):
    """Aggiungi entry al log di processing"""
    entry["timestamp"] = datetime.now().timestamp()
    rc.lpush(Q_REDIS_KEY_PROCESSING_LOG, json.dumps(entry))
    rc.ltrim(Q_REDIS_KEY_PROCESSING_LOG, 0, 99)

def _update_stats(rc: Redis, **kwargs):
    """Aggiorna statistiche aggregate"""
    stats_raw = rc.get(Q_REDIS_KEY_STATS)
    if stats_raw:
        stats = json.loads(stats_raw)
    else:
        stats = {
            "success": 0,
            "failed": 0,
            "chunked": 0,
            "meili_indexed": 0,
            "qdrant_vectorized": 0
        }
    
    stats.update(kwargs)
    rc.set(Q_REDIS_KEY_STATS, json.dumps(stats))

def _clean_text(text: str) -> str:
    """Pulisce testo da caratteri problematici per PostgreSQL"""
    if not text:
        return ""
    
    text = text.replace('\x00', '').replace('\0', '')
    
    import re
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def _pdftotext_safe(path: str) -> str:
    """Estrazione PDF con gestione errori"""
    try:
        out = subprocess.run(
            ["pdftotext", "-layout", "-nopgbrk", path, "-"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=30
        )
        
        if out.returncode != 0:
            log.warning(f"pdftotext fallito per {os.path.basename(path)}")
            return ""
        
        return _clean_text(out.stdout)
    
    except subprocess.TimeoutExpired:
        log.error(f"pdftotext timeout su {os.path.basename(path)}")
        return ""
    except Exception as e:
        log.error(f"pdftotext errore su {os.path.basename(path)}: {e}")
        return ""

def _extract_docx(path: str) -> str:
    """Estrazione DOCX con fallback a LibreOffice"""
    try:
        from docx import Document
        
        doc = Document(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        all_text = "\n".join(paragraphs + tables_text)
        return _clean_text(all_text)
    
    except ImportError:
        log.debug(f"python-docx non disponibile, uso LibreOffice per {os.path.basename(path)}")
        return _libreoffice_convert_safe(path)
    
    except Exception as e:
        log.warning(f"python-docx fallito su {os.path.basename(path)}: {e}")
        return _libreoffice_convert_safe(path)

def _libreoffice_convert_safe(path: str) -> str:
    """Conversione LibreOffice con gestione errori"""
    import tempfile
    
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "txt:Text",
                "--outdir", tmpdir,
                path
            ],
            capture_output=True,
            text=True,
            timeout=60)
            
            if result.returncode != 0:
                log.warning(f"LibreOffice fallito su {os.path.basename(path)}")
                return ""
            
            basename = os.path.splitext(os.path.basename(path))[0]
            txt_files = [f for f in os.listdir(tmpdir) if f.endswith('.txt')]
            
            if not txt_files:
                return ""
            
            txt_file = None
            for f in txt_files:
                if basename.lower() in f.lower():
                    txt_file = f
                    break
            
            if not txt_file:
                txt_file = txt_files[0]
            
            txt_path = os.path.join(tmpdir, txt_file)
            
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            
            return _clean_text(content)
    
    except subprocess.TimeoutExpired:
        log.error(f"LibreOffice timeout su {os.path.basename(path)}")
        return ""
    
    except Exception as e:
        log.error(f"LibreOffice errore su {os.path.basename(path)}: {e}")
        return ""

def _read_text(path: str) -> str:
    """Estrae testo da vari formati"""
    ext = os.path.splitext(path)[1].lower()
    
    UNSUPPORTED = {
        '.mpp', '.vsd', '.mdb', '.accdb',
        '.zip', '.rar', '.7z', '.tar', '.gz',
        '.exe', '.dll', '.so',
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg',
        '.mp3', '.mp4', '.avi', '.mov', '.wav',
    }
    
    if ext in UNSUPPORTED:
        return ""
    
    try:
        if ext in (".txt", ".md", ".csv", ".log", ".ini", ".conf", ".xml", ".json"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return _clean_text(f.read())
        
        if ext == ".pdf":
            return _pdftotext_safe(path)
        
        if ext == ".docx":
            return _extract_docx(path)
        
        if ext in (".doc", ".xls", ".xlsx", ".ppt", ".pptx", ".odt", ".ods", ".odp", ".rtf"):
            return _libreoffice_convert_safe(path)
        
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _clean_text(f.read())
    
    except Exception as e:
        log.error(f"Errore estrazione {os.path.basename(path)}: {e}")
        return ""

def _collect_files(root: str) -> List[str]:
    """Raccoglie tutti i file dalla directory root"""
    files = []
    for base, _, names in os.walk(root):
        for n in names:
            if n.startswith("."):
                continue
            p = os.path.join(base, n)
            if os.path.isfile(p):
                files.append(p)
    return files

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Divide testo in chunk con overlap"""
    if not text or len(text) < chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            cut_point = max(last_period, last_newline)
            
            if cut_point > chunk_size // 2:
                chunk = chunk[:cut_point + 1]
                end = start + cut_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap if end < len(text) else end
    
    return chunks

def _get_embedder(model_type: str):
    """Ottieni embedder per il modello specificato CON GPU SUPPORT"""
    config = MODEL_CONFIGS.get(model_type)
    if not config:
        raise ValueError(f"Modello non supportato: {model_type}")
    
    if config["type"] == "transformers":
        from sentence_transformers import SentenceTransformer
        
        log.info(f"🔄 Worker caricamento modello {config['name']} su {DEVICE}...")
        
        # Carica modello su device appropriato
        model = SentenceTransformer(config["name"], device=DEVICE)
        
        log.info(f"✅ Worker modello {config['name']} caricato su {DEVICE}")
        
        def embed_batch(texts: List[str]) -> List[List[float]]:
            """Embedding batch con GPU support e fallback CPU"""
            try:
                embeddings = model.encode(
                    texts,
                    convert_to_numpy=True,
                    batch_size=BATCH_SIZE,
                    show_progress_bar=False,
                    device=DEVICE
                )
                
                # Clear cache periodicamente
                if GPU_AVAILABLE and len(texts) > 32:
                    clear_gpu_cache()
                
                return embeddings.tolist()
            
            except Exception as e:
                log.error(f"Errore embedding batch: {e}")
                
                # Fallback CPU se GPU OOM
                if GPU_AVAILABLE and "out of memory" in str(e).lower():
                    log.warning("⚠️ GPU OOM, retry con CPU...")
                    clear_gpu_cache()
                    
                    embeddings = model.encode(
                        texts,
                        convert_to_numpy=True,
                        batch_size=16,
                        show_progress_bar=False,
                        device='cpu'
                    )
                    return embeddings.tolist()
                
                raise
        
        return embed_batch
    
    elif config["type"] == "ollama":
        import requests
        
        def embed_ollama(texts: List[str]) -> List[List[float]]:
            """Embeddings via Ollama API"""
            embeddings = []
            for text in texts:
                try:
                    resp = requests.post(
                        "http://ollama:11434/api/embeddings",
                        json={"model": config["name"], "prompt": text},
                        timeout=30
                    )
                    if resp.status_code == 200:
                        embeddings.append(resp.json()["embedding"])
                    else:
                        log.error(f"Ollama errore: {resp.status_code}")
                        embeddings.append([0.0] * config["dimension"])
                except Exception as e:
                    log.error(f"Ollama errore: {e}")
                    embeddings.append([0.0] * config["dimension"])
            return embeddings
        
        return embed_ollama
    
    raise ValueError(f"Tipo modello non supportato: {config['type']}")

def extract_metadata(filepath: str, kb_root: str = KB_ROOT) -> Dict[str, Optional[str]]:
    """
    Estrae metadati strutturati da path e nome file.
    """
    import re
    from pathlib import Path
    
    metadata = {
        "area": None,
        "anno": None,
        "cliente": None,
        "oggetto": None,
        "tipo_doc": None,
        "codice_appalto": None,
        "categoria": None,
        "versione": None,
        "ext": None
    }
    
    rel_path = filepath.replace(kb_root, "").lstrip("/")
    parts = rel_path.split("/")
    
    if len(parts) < 2:
        return metadata
    
    _, ext = os.path.splitext(filepath)
    metadata["ext"] = ext.lstrip(".").lower()
    
    area_folder = parts[0]
    if area_folder.startswith("_"):
        metadata["area"] = area_folder.lstrip("_")
    
    if metadata["area"] == "AQ":
        STRALCIO_MAP = {"SD1": "2021", "SD2": "2022", "SD3": "2023", "SD4": "2024", "SD5": "2025", "SD6": "2026"}
        
        if len(parts) > 1 and parts[1].startswith("SD"):
            metadata["anno"] = STRALCIO_MAP.get(parts[1])
        
        for part in parts:
            match = re.search(r'\b(AS\d{4}[_A-Z0-9]*)\b', part)
            if match:
                metadata["codice_appalto"] = match.group(1)
                if "_" in part:
                    possible_cliente = part.split("_")[1:]
                    if possible_cliente:
                        metadata["cliente"] = "_".join(possible_cliente)
                break
        
        TIPO_DOC_ALIASES = {
            "01_Documentazione": "Documentazione",
            "02_Chiarimenti": "Chiarimenti",
            "04_OffertaTecnica": "Offerta Tecnica",
            "08_AccessoAgliAtti": "Accesso Atti",
            "99_AS": "Appalto Specifico"
        }
        
        for part in parts:
            if re.match(r'^\d{2}_', part):
                metadata["tipo_doc"] = TIPO_DOC_ALIASES.get(part, part)
                break
    
    elif metadata["area"] == "Gare":
        gara_folder = parts[1]
        
        match = re.match(r'^(\d{4})_(.+?)(?:-(.+))?$', gara_folder)
        if match:
            metadata["anno"] = match.group(1)
            cliente_raw = match.group(2)
            oggetto_raw = match.group(3) if match.group(3) else ""
            
            metadata["cliente"] = re.sub(r'^(AOU|AO|ASL|AUSL|ASP|Regione|Provincia)\s*', '', cliente_raw, flags=re.IGNORECASE).strip()
            
            if oggetto_raw:
                metadata["oggetto"] = oggetto_raw.replace("_", " ")
    
    filename = os.path.basename(filepath)
    version_match = re.search(r'[vV]\.?\d+\.\d+(?:\.\d+)?', filename)
    if version_match:
        metadata["versione"] = version_match.group(0)
    
    return metadata

# ===== MAIN TASK =====
def run_ingestion(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Task principale di ingestion multi-modello CON GPU ACCELERATION.
    
    Args:
        params: {
            "mode": "full" | "incremental",
            "model": "sentence-transformer" | "llama3" | "mistral",
            "use_gpu": True | False (opzionale)
        }
    """
    mode = params.get("mode", "full")
    model_type = params.get("model", "sentence-transformer")
    
    if model_type not in MODEL_CONFIGS:
        return {"ok": False, "error": f"Modello non supportato: {model_type}"}
    
    config = MODEL_CONFIGS[model_type]
    collection_name = f"{config['collection_prefix']}_docs"
    
    log.info(f"🚀 Avvio ingestion: mode={mode}, model={model_type}, collection={collection_name}, device={DEVICE}")
    
    rc = rconn()
    ensure_pg_schema()
    
    rc.delete(Q_REDIS_KEY_FAILED)
    rc.delete(Q_REDIS_KEY_CURRENT_DOC)
    rc.set(Q_REDIS_KEY_STATS, json.dumps({
        "success": 0, "failed": 0, "chunked": 0, "meili_indexed": 0, "qdrant_vectorized": 0
    }))
    
    all_files = _collect_files(KB_ROOT)
    total = len(all_files)
    done = 0
    
    _set_progress(rc, True, 0, total, f"init-{model_type}")
    
    pg = pg_conn()
    meili = meili_client()
    qd = qdrant_client()
    
    # Crea collection Qdrant
    try:
        collections = qd.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            qd.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=config["dimension"], distance=Distance.COSINE)
            )
            log.info(f"✅ Collection Qdrant creata: {collection_name}")
    except Exception as e:
        log.error(f"Errore Qdrant init: {e}")
        _push_failed(rc, {"stage": "qdrant-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    # Init Meilisearch
    try:
        try:
            meili.get_index(MEILI_INDEX)
        except meilisearch.errors.MeilisearchApiError:
            meili.create_index(MEILI_INDEX, {"primaryKey": "id"})
        
        idx = meili.index(MEILI_INDEX)
        idx.update_filterable_attributes([
            "area", "anno", "cliente", "oggetto", "tipo_doc", "categoria", "ext"
        ])
    except Exception as e:
        log.error(f"Errore Meilisearch init: {e}")
        _push_failed(rc, {"stage": "meili-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    # Get embedder CON GPU
    try:
        embedder = _get_embedder(model_type)
    except Exception as e:
        log.error(f"Errore caricamento modello {model_type}: {e}")
        _push_failed(rc, {"stage": "model-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    meili_batch = []
    total_chunks = 0
    
    with pg, pg.cursor() as cur:
        for path in all_files:
            # Check pause
            if rc.exists("kb:ingestion_pause"):
                log.info("⏸️ Ingestion in pausa")
                _set_progress(rc, False, done, total, "paused")
                while rc.exists("kb:ingestion_pause"):
                    import time
                    time.sleep(2)
                log.info("▶️ Ingestion ripresa")
                _set_progress(rc, True, done, total, f"processing-{model_type}")
            
            rel_id = os.path.relpath(path, KB_ROOT)
            filename = os.path.basename(path)
            
            _set_current_doc(rc, {
                "filename": filename,
                "path": path,
                "step": "reading"
            })
            
            try:
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "extracting"})
                text = _read_text(path)
                
                if not text:
                    log.warning(f"⚠️ Nessun testo estratto da {filename}")
                    done += 1
                    continue
                
                title = filename
                metadata = extract_metadata(path, KB_ROOT)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "postgres"})
                cur.execute("""
                    INSERT INTO documents (
                        id, path, title, content,
                        ext, area, anno, cliente, oggetto, tipo_doc,
                        codice_appalto, categoria, versione
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET
                        path=EXCLUDED.path,
                        title=EXCLUDED.title,
                        content=EXCLUDED.content,
                        ext=EXCLUDED.ext,
                        area=EXCLUDED.area,
                        anno=EXCLUDED.anno,
                        cliente=EXCLUDED.cliente,
                        oggetto=EXCLUDED.oggetto,
                        tipo_doc=EXCLUDED.tipo_doc,
                        codice_appalto=EXCLUDED.codice_appalto,
                        categoria=EXCLUDED.categoria,
                        versione=EXCLUDED.versione,
                        mtime=NOW()
                """, (
                    rel_id, path, title, text,
                    metadata.get('ext'),
                    metadata.get('area'),
                    metadata.get('anno'),
                    metadata.get('cliente'),
                    metadata.get('oggetto'),
                    metadata.get('tipo_doc'),
                    metadata.get('codice_appalto'),
                    metadata.get('categoria'),
                    metadata.get('versione')
                ))
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "chunking"})
                chunks = _chunk_text(text, chunk_size=1500, overlap=200)
                
                if not chunks:
                    done += 1
                    continue
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "embedding", "details": f"{len(chunks)} chunks ({DEVICE})"})
                embeddings = embedder(chunks)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "qdrant"})
                points = []
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    point_id = _generate_point_id(rel_id, i)
                    
                    points.append(PointStruct(
                        id=point_id,
                        vector=embedding,
                        payload={
                            "doc_id": rel_id,
                            "chunk_id": i,
                            "text": chunk,
                            "title": title,
                            "path": path,
                            **{k: v for k, v in metadata.items() if v is not None}
                        }
                    ))
                
                qd.upsert(collection_name=collection_name, points=points)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "meili"})
                meili_batch.append({
                    "id": rel_id,
                    "path": path,
                    "title": title,
                    "content": text[:5000],
                    **{k: v for k, v in metadata.items() if v is not None}
                })
                
                total_chunks += len(chunks)
                
                _add_processing_log(rc, {
                    "filename": filename,
                    "status": "success",
                    "chunks": len(chunks),
                    "meili_indexed": True,
                    "qdrant_vectorized": True,
                    "device": DEVICE
                })
                
                _update_stats(rc,
                    success=done + 1,
                    chunked=total_chunks,
                    qdrant_vectorized=done + 1
                )
            
            except Exception as e:
                log.error(f"❌ Errore processing {filename}: {e}")
                import traceback
                log.error(traceback.format_exc())
                
                _push_failed(rc, {"path": path, "filename": filename, "error": str(e)})
                _add_processing_log(rc, {
                    "filename": filename,
                    "status": "error",
                    "error": str(e)
                })
                _update_stats(rc, failed=rc.llen(Q_REDIS_KEY_FAILED))
            
            finally:
                done += 1
                
                if len(meili_batch) >= 50:
                    try:
                        idx.add_documents(meili_batch)
                        _update_stats(rc, meili_indexed=done)
                    except Exception as e:
                        log.error(f"Errore batch Meilisearch: {e}")
                    meili_batch = []
                
                _set_progress(rc, True, done, total, f"processing-{model_type}")
        
        if meili_batch:
            try:
                idx.add_documents(meili_batch)
                _update_stats(rc, meili_indexed=done)
            except Exception as e:
                log.error(f"Errore final batch: {e}")
    
    _set_current_doc(rc, None)
    _set_progress(rc, False, total, total, "done")
    
    log.info(f"✅ Ingestion completata: {total} documenti, {total_chunks} chunks, modello {model_type}, device {DEVICE}")
    
    # Clear GPU cache finale
    if GPU_AVAILABLE:
        clear_gpu_cache()
    
    return {
        "ok": True,
        "mode": mode,
        "model": model_type,
        "collection": collection_name,
        "total": total,
        "chunks": total_chunks,
        "device": DEVICE
    }
