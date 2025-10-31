# worker/worker_tasks.py - IMPROVED VERSION
# Cambiamenti rispetto all'originale:
# 1. Timeout dinamico basato su dimensione file
# 2. Blacklist estesa per file problematici
# 3. Gestione più robusta dei processi LibreOffice
# 4. Fallback strategy

import os
import json
import subprocess
import logging
import uuid
import hashlib
import signal
from typing import Dict, Any, List, Optional
from datetime import datetime

import redis
from redis import Redis

import psycopg
from psycopg.rows import dict_row

import meilisearch
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

# ===== GPU DETECTION =====
import torch

DEVICE = None
GPU_AVAILABLE = False

try:
    if torch.cuda.is_available():
        DEVICE = 'cuda'
        GPU_AVAILABLE = True
        logging.info(f"🎮 Worker GPU DETECTED: {torch.cuda.get_device_name(0)}")
        logging.info(f"🎮 CUDA Version: {torch.version.cuda}")
    else:
        DEVICE = 'cpu'
        logging.info("💻 GPU non disponibile, worker usa CPU")
except Exception as e:
    DEVICE = 'cpu'
    logging.warning(f"⚠️ Errore detection GPU: {e}, worker usa CPU")

# Batch size ottimizzato per device
BATCH_SIZE = 48 if GPU_AVAILABLE else 16
logging.info(f"📦 Worker batch size: {BATCH_SIZE} ({'GPU' if GPU_AVAILABLE else 'CPU'} optimized)")

# ===== ENV =====
REDIS_URL = os.getenv("REDIS_URL", "redis://redis:6379/0")
RQ_QUEUE = os.getenv("RQ_QUEUE", "kb_ingestion")

POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres")
POSTGRES_DB = os.getenv("POSTGRES_DB", "kb")
POSTGRES_USER = os.getenv("POSTGRES_USER", "kbuser")
POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD", "kbpass")

MEILI_URL = os.getenv("MEILI_URL", "http://meili:7700")
MEILI_MASTER_KEY = os.getenv("MEILI_MASTER_KEY", os.getenv("MEILI_KEY", "change_me_meili_key"))

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

KB_ROOT = os.getenv("KB_ROOT", "/mnt/kb")
MEILI_INDEX = "kb_docs"

Q_REDIS_KEY_PROGRESS = "kb:progress"
Q_REDIS_KEY_FAILED = "kb:failed_docs"
Q_REDIS_KEY_CURRENT_DOC = "kb:current_doc"
Q_REDIS_KEY_PROCESSING_LOG = "kb:processing_log"
Q_REDIS_KEY_STATS = "kb:stats"

# Configurazione modelli
MODEL_CONFIGS = {
    "sentence-transformer": {
        "name": "all-MiniLM-L6-v2",
        "dimension": 384,
        "collection_prefix": "kb_st",
        "type": "transformers"
    },
    "llama3": {
        "name": "llama3",
        "dimension": 4096,
        "collection_prefix": "kb_llama3",
        "type": "ollama"
    },
    "mistral": {
        "name": "mistral",
        "dimension": 4096,
        "collection_prefix": "kb_mistral",
        "type": "ollama"
    }
}

log = logging.getLogger("worker")

# ===== GPU Memory Management =====
def clear_gpu_cache():
    """Libera memoria GPU cache"""
    if GPU_AVAILABLE:
        try:
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        except Exception as e:
            log.warning(f"Errore clear GPU cache: {e}")

# ===== Helpers =====
def rconn() -> Redis:
    return redis.from_url(REDIS_URL, decode_responses=True)

def meili_client() -> meilisearch.Client:
    return meilisearch.Client(MEILI_URL, MEILI_MASTER_KEY)

def qdrant_client() -> QdrantClient:
    return QdrantClient(url=QDRANT_URL)

def pg_conn():
    dsn = f"host={POSTGRES_HOST} dbname={POSTGRES_DB} user={POSTGRES_USER} password={POSTGRES_PASSWORD}"
    return psycopg.connect(dsn, autocommit=True, row_factory=dict_row)

def ensure_pg_schema():
    """Crea schema PostgreSQL con metadati avanzati"""
    with pg_conn() as conn, conn.cursor() as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY,
            path TEXT NOT NULL,
            title TEXT,
            content TEXT,
            mtime TIMESTAMP DEFAULT NOW(),
            
            -- Metadati base
            ext TEXT,
            
            -- Metadati strutturati (da metadata_extractor)
            area TEXT,
            anno TEXT,
            cliente TEXT,
            oggetto TEXT,
            tipo_doc TEXT,
            codice_appalto TEXT,
            categoria TEXT,
            descrizione_oggetto TEXT,
            versione TEXT
        );
        """)
        
        # Indici per performance
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_area ON documents(area);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_anno ON documents(anno);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_cliente ON documents(cliente);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_oggetto ON documents(oggetto);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_categoria ON documents(categoria);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_docs_ext ON documents(ext);")

def _generate_point_id(doc_id: str, chunk_idx: int) -> str:
    """
    Genera UUID deterministico per point ID Qdrant.
    Usa hash MD5 per garantire stesso ID per stesso documento+chunk.
    """
    unique_str = f"{doc_id}_chunk_{chunk_idx}"
    namespace = uuid.UUID('6ba7b810-9dad-11d1-80b4-00c04fd430c8')
    point_uuid = uuid.uuid5(namespace, unique_str)
    return str(point_uuid)

def _push_failed(rc: Redis, item: Dict[str, Any]):
    """Aggiungi documento fallito alla lista"""
    rc.lpush(Q_REDIS_KEY_FAILED, json.dumps({**item, "ts": datetime.now().timestamp()}))

def _set_progress(rc: Redis, running: bool, done: int, total: int, stage: str):
    """Aggiorna progress generale"""
    rc.set(Q_REDIS_KEY_PROGRESS, json.dumps({
        "running": running, "done": done, "total": total, "stage": stage
    }))

def _set_current_doc(rc: Redis, doc_info: Optional[Dict[str, Any]]):
    """Imposta documento corrente in elaborazione"""
    if doc_info:
        rc.set(Q_REDIS_KEY_CURRENT_DOC, json.dumps({**doc_info, "ts": datetime.now().timestamp()}))
    else:
        rc.delete(Q_REDIS_KEY_CURRENT_DOC)

def _add_processing_log(rc: Redis, entry: Dict[str, Any]):
    """Aggiungi entry al log di processing"""
    entry["timestamp"] = datetime.now().timestamp()
    rc.lpush(Q_REDIS_KEY_PROCESSING_LOG, json.dumps(entry))
    rc.ltrim(Q_REDIS_KEY_PROCESSING_LOG, 0, 99)

def _update_stats(rc: Redis, **kwargs):
    """Aggiorna statistiche aggregate"""
    stats_raw = rc.get(Q_REDIS_KEY_STATS)
    if stats_raw:
        stats = json.loads(stats_raw)
    else:
        stats = {
            "success": 0,
            "failed": 0,
            "chunked": 0,
            "meili_indexed": 0,
            "qdrant_vectorized": 0
        }
    
    stats.update(kwargs)
    rc.set(Q_REDIS_KEY_STATS, json.dumps(stats))

def _clean_text(text: str) -> str:
    """Pulisce testo da caratteri problematici per PostgreSQL"""
    if not text:
        return ""
    
    text = text.replace('\x00', '').replace('\0', '')
    
    import re
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def _get_file_size_mb(path: str) -> float:
    """Ottiene dimensione file in MB"""
    try:
        return os.path.getsize(path) / (1024 * 1024)
    except:
        return 0

def _calculate_timeout(path: str, base_timeout: int = 30) -> int:
    """
    Calcola timeout dinamico basato su dimensione file.
    
    Strategia:
    - File <5MB: 30s (default, la maggior parte dei file)
    - File 5-20MB: 60s (documenti grandi ma gestibili)
    - File 20-50MB: 90s (documenti molto grandi)
    - File >50MB: 120s (massimo, file eccezionali)
    
    Questo permette di:
    1. Processare velocemente file piccoli
    2. Dare tempo sufficiente a file legittimi grandi
    3. Non bloccarsi troppo su file problematici
    """
    size_mb = _get_file_size_mb(path)
    
    if size_mb < 5:
        return 30
    elif size_mb < 20:
        return 60
    elif size_mb < 50:
        return 90
    else:
        # File molto grandi: max 120s, poi skip
        log.warning(f"File molto grande ({size_mb:.1f}MB): {os.path.basename(path)}")
        return 120

def _pdftotext_safe(path: str) -> str:
    """Estrazione PDF con gestione errori e timeout dinamico"""
    timeout = _calculate_timeout(path, base_timeout=30)
    
    try:
        out = subprocess.run(
            ["pdftotext", "-layout", "-nopgbrk", path, "-"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout
        )
        
        if out.returncode != 0:
            log.warning(f"pdftotext fallito per {os.path.basename(path)}")
            return ""
        
        return _clean_text(out.stdout)
    
    except subprocess.TimeoutExpired:
        log.error(f"⏱️ pdftotext timeout ({timeout}s) su {os.path.basename(path)} ({_get_file_size_mb(path):.1f}MB)")
        return ""
    except Exception as e:
        log.error(f"pdftotext errore su {os.path.basename(path)}: {e}")
        return ""

def _extract_docx(path: str) -> str:
    """Estrazione DOCX con fallback a LibreOffice"""
    try:
        from docx import Document
        
        doc = Document(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        all_text = "\n".join(paragraphs + tables_text)
        return _clean_text(all_text)
    
    except ImportError:
        log.debug(f"python-docx non disponibile, uso LibreOffice per {os.path.basename(path)}")
        return _libreoffice_convert_safe(path)
    
    except Exception as e:
        log.warning(f"python-docx fallito su {os.path.basename(path)}: {e}")
        return _libreoffice_convert_safe(path)

def _kill_libreoffice_for_file(path: str):
    """Tenta di killare processo LibreOffice specifico per un file"""
    try:
        basename = os.path.basename(path)
        # Cerca processi libreoffice che contengono il nome del file
        result = subprocess.run(
            ["pgrep", "-f", basename],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0 and result.stdout.strip():
            pids = result.stdout.strip().split('\n')
            for pid in pids:
                try:
                    os.kill(int(pid), signal.SIGKILL)
                    log.info(f"Killed process {pid} for {basename}")
                except:
                    pass
    except Exception as e:
        log.debug(f"Errore kill specifico: {e}")

def _libreoffice_convert_safe(path: str) -> str:
    """
    Conversione LibreOffice con timeout dinamico e gestione robusta.
    
    Miglioramenti:
    1. Timeout basato su dimensione file
    2. Kill processo specifico dopo timeout
    3. Cleanup tempdir garantito
    4. Log dettagliato con dimensioni file
    """
    import tempfile
    
    # Calcola timeout dinamico
    timeout = _calculate_timeout(path)
    size_mb = _get_file_size_mb(path)
    filename = os.path.basename(path)
    
    log.debug(f"LibreOffice: {filename} ({size_mb:.1f}MB, timeout={timeout}s)")
    
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "txt:Text",
                "--outdir", tmpdir,
                path
            ],
            capture_output=True,
            text=True,
            timeout=timeout)
            
            if result.returncode != 0:
                log.warning(f"LibreOffice fallito su {filename}: returncode {result.returncode}")
                if result.stderr:
                    log.debug(f"stderr: {result.stderr[:200]}")
                return ""
            
            basename = os.path.splitext(filename)[0]
            txt_files = [f for f in os.listdir(tmpdir) if f.endswith('.txt')]
            
            if not txt_files:
                log.warning(f"LibreOffice non ha prodotto output per {filename}")
                return ""
            
            txt_file = None
            for f in txt_files:
                if basename.lower() in f.lower():
                    txt_file = f
                    break
            
            if not txt_file:
                txt_file = txt_files[0]
            
            txt_path = os.path.join(tmpdir, txt_file)
            
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            
            if content:
                log.debug(f"✅ LibreOffice OK: {filename} → {len(content)} chars")
            
            return _clean_text(content)
    
    except subprocess.TimeoutExpired:
        log.error(f"⏱️ LibreOffice timeout ({timeout}s) su {filename} ({size_mb:.1f}MB)")
        
        # Tenta kill specifico del processo
        _kill_libreoffice_for_file(path)
        
        return ""
    
    except Exception as e:
        log.error(f"LibreOffice errore su {filename}: {e}")
        return ""

def _read_text(path: str) -> str:
    """
    Estrae testo da vari formati con blacklist estesa.
    
    Miglioramenti:
    1. Blacklist più completa con file CAD, Project, ecc.
    2. Gestione esplicita per ogni formato
    3. Fallback controllato
    """
    ext = os.path.splitext(path)[1].lower()
    
    # ===== BLACKLIST ESTESA =====
    # File che NON devono essere processati perché:
    # - Non supportati da LibreOffice
    # - Causano crash/hang
    # - Non contengono testo estraibile
    UNSUPPORTED = {
        # CAD e Design
        '.dwg', '.dxf', '.dwf',          # AutoCAD
        '.skp',                           # SketchUp
        
        # Project Management
        '.mpp',                           # Microsoft Project
        
        # Diagrammi
        '.vsd', '.vsdx',                  # Microsoft Visio
        
        # Database
        '.mdb', '.accdb',                 # Microsoft Access
        '.db', '.sqlite', '.sqlite3',     # SQLite
        
        # Archivi
        '.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz',
        
        # Eseguibili e Libraries
        '.exe', '.dll', '.so', '.dylib', '.app',
        
        # Immagini (non OCR per ora)
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.ico', '.webp',
        
        # Audio/Video
        '.mp3', '.mp4', '.avi', '.mov', '.wav', '.flac', '.mkv', '.wmv',
        
        # Outlook
        '.pst', '.ost',                   # Outlook data files
        
        # Font
        '.ttf', '.otf', '.woff', '.woff2',
    }
    
    if ext in UNSUPPORTED:
        log.debug(f"Skip {ext} (blacklist): {os.path.basename(path)}")
        return ""
    
    try:
        # Text files - lettura diretta
        if ext in (".txt", ".md", ".csv", ".log", ".ini", ".conf", ".xml", ".json", ".yaml", ".yml"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return _clean_text(f.read())
        
        # PDF - tool dedicato
        if ext == ".pdf":
            return _pdftotext_safe(path)
        
        # DOCX - python-docx con fallback LibreOffice
        if ext == ".docx":
            return _extract_docx(path)
        
        # Altri formati Office - LibreOffice con timeout dinamico
        if ext in (".doc", ".xls", ".xlsx", ".ppt", ".pptx", ".odt", ".ods", ".odp", ".rtf"):
            return _libreoffice_convert_safe(path)
        
        # Fallback: prova lettura come testo
        # (es: .py, .js, .cpp, ecc.)
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                # Se sembra testo (>80% ASCII stampabile), ritorna
                if len(content) > 0:
                    printable = sum(c.isprintable() or c.isspace() for c in content[:1000])
                    if printable / len(content[:1000]) > 0.8:
                        return _clean_text(content)
        except:
            pass
        
        # Se arriviamo qui, non sappiamo come gestirlo
        log.debug(f"Skip {ext} (no handler): {os.path.basename(path)}")
        return ""
    
    except Exception as e:
        log.error(f"Errore estrazione {os.path.basename(path)}: {e}")
        return ""

# ===== IL RESTO DEL FILE RIMANE IDENTICO =====
# (Le funzioni _collect_files, _chunk_text, _get_embedder, extract_metadata, run_ingestion
#  rimangono identiche all'originale)

def _collect_files(root: str) -> List[str]:
    """Raccoglie tutti i file dalla directory root"""
    files = []
    for base, _, names in os.walk(root):
        for n in names:
            if n.startswith("."):
                continue
            p = os.path.join(base, n)
            if os.path.isfile(p):
                files.append(p)
    return files

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Divide testo in chunk con overlap"""
    if not text or len(text) < chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            cut_point = max(last_period, last_newline)
            
            if cut_point > chunk_size // 2:
                chunk = chunk[:cut_point + 1]
                end = start + cut_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap if end < len(text) else end
    
    return chunks

def _get_embedder(model_type: str):
    """Ottieni embedder per il modello specificato CON GPU SUPPORT"""
    config = MODEL_CONFIGS.get(model_type)
    if not config:
        raise ValueError(f"Modello non supportato: {model_type}")
    
    if config["type"] == "transformers":
        from sentence_transformers import SentenceTransformer
        
        log.info(f"🔄 Worker caricamento modello {config['name']} su {DEVICE}...")
        
        # Carica modello su device appropriato
        model = SentenceTransformer(config["name"], device=DEVICE)
        
        log.info(f"✅ Modello {config['name']} caricato su {DEVICE}")
        
        def embed_batch(texts: List[str]) -> List[List[float]]:
            """Embedding batch con GPU"""
            try:
                # encode() usa automaticamente il device del modello
                embeddings = model.encode(
                    texts,
                    batch_size=BATCH_SIZE,
                    show_progress_bar=False,
                    convert_to_numpy=True
                )
                return embeddings.tolist()
            except Exception as e:
                log.error(f"Errore embedding batch: {e}")
                raise
        
        return embed_batch
    
    elif config["type"] == "ollama":
        import requests
        
        OLLAMA_URL = os.getenv("OLLAMA_URL", "http://ollama:11434")
        
        def embed_batch(texts: List[str]) -> List[List[float]]:
            embeddings = []
            for text in texts:
                resp = requests.post(
                    f"{OLLAMA_URL}/api/embeddings",
                    json={"model": config["name"], "prompt": text},
                    timeout=30
                )
                resp.raise_for_status()
                embeddings.append(resp.json()["embedding"])
            return embeddings
        
        return embed_batch
    
    else:
        raise ValueError(f"Tipo modello non supportato: {config['type']}")

def extract_metadata(file_path: str, kb_root: str) -> Dict[str, Any]:
    """Estrae metadati strutturati dal path"""
    rel_path = os.path.relpath(file_path, kb_root)
    parts = rel_path.split(os.sep)
    
    metadata = {
        'ext': os.path.splitext(file_path)[1].lower()
    }
    
    if len(parts) >= 2 and parts[0].startswith('_'):
        metadata['area'] = parts[0][1:]
    
    if len(parts) >= 2:
        folder_name = parts[1] if parts[0].startswith('_') else parts[0]
        
        import re
        year_match = re.search(r'(19|20)\d{2}', folder_name)
        if year_match:
            metadata['anno'] = year_match.group(0)
        
        if '-' in folder_name:
            parts_folder = folder_name.split('-', 1)
            if len(parts_folder) == 2:
                metadata['cliente'] = parts_folder[0].strip()
                metadata['oggetto'] = parts_folder[1].strip()
    
    filename = os.path.basename(file_path)
    lower_filename = filename.lower()
    
    if 'offerta' in lower_filename or 'offer' in lower_filename:
        metadata['tipo_doc'] = 'offerta'
    elif 'capitolato' in lower_filename:
        metadata['tipo_doc'] = 'capitolato'
    elif 'contratto' in lower_filename:
        metadata['tipo_doc'] = 'contratto'
    elif 'relazione' in lower_filename:
        metadata['tipo_doc'] = 'relazione'
    elif 'manuale' in lower_filename:
        metadata['tipo_doc'] = 'manuale'
    
    return metadata

def run_ingestion(mode: str = "full", model_type: str = "sentence-transformer") -> Dict[str, Any]:
    """
    Esegue ingestion con GPU support e timeout dinamici.
    
    mode:
        - "full": tutto KB_ROOT
        - "gare": solo _Gare
        - "aq": solo _AQ
    
    model_type:
        - "sentence-transformer": all-MiniLM-L6-v2 (384d)
        - "llama3": llama3 via Ollama (4096d)
        - "mistral": mistral via Ollama (4096d)
    """
    
    config = MODEL_CONFIGS.get(model_type)
    if not config:
        return {"ok": False, "error": f"Modello non supportato: {model_type}"}
    
    collection_name = f"{config['collection_prefix']}_chunks"
    
    log.info(f"🚀 Avvio ingestion mode={mode}, model={model_type}, collection={collection_name}, device={DEVICE}")
    
    ensure_pg_schema()
    
    if mode == "full":
        root = KB_ROOT
    elif mode == "gare":
        root = os.path.join(KB_ROOT, "_Gare")
    elif mode == "aq":
        root = os.path.join(KB_ROOT, "_AQ")
    else:
        return {"ok": False, "error": f"Mode non valido: {mode}"}
    
    all_files = _collect_files(root)
    total = len(all_files)
    done = 0
    
    log.info(f"📂 Trovati {total} file in {root}")
    
    rc = rconn()
    _set_progress(rc, True, 0, total, f"init-{model_type}")
    
    pg = pg_conn()
    meili = meili_client()
    qd = qdrant_client()
    
    # Crea collection Qdrant
    try:
        collections = qd.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            qd.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=config["dimension"], distance=Distance.COSINE)
            )
            log.info(f"✅ Collection Qdrant creata: {collection_name}")
    except Exception as e:
        log.error(f"Errore Qdrant init: {e}")
        _push_failed(rc, {"stage": "qdrant-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    # Init Meilisearch
    try:
        try:
            meili.get_index(MEILI_INDEX)
        except meilisearch.errors.MeilisearchApiError:
            meili.create_index(MEILI_INDEX, {"primaryKey": "id"})
        
        idx = meili.index(MEILI_INDEX)
        idx.update_filterable_attributes([
            "area", "anno", "cliente", "oggetto", "tipo_doc", "categoria", "ext"
        ])
    except Exception as e:
        log.error(f"Errore Meilisearch init: {e}")
        _push_failed(rc, {"stage": "meili-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    # Get embedder CON GPU
    try:
        embedder = _get_embedder(model_type)
    except Exception as e:
        log.error(f"Errore caricamento modello {model_type}: {e}")
        _push_failed(rc, {"stage": "model-init", "error": str(e)})
        _set_progress(rc, False, 0, 0, "failed")
        return {"ok": False, "error": str(e)}
    
    meili_batch = []
    total_chunks = 0
    
    with pg, pg.cursor() as cur:
        for path in all_files:
            # Check pause
            if rc.exists("kb:ingestion_pause"):
                log.info("⏸️ Ingestion in pausa")
                _set_progress(rc, False, done, total, "paused")
                while rc.exists("kb:ingestion_pause"):
                    import time
                    time.sleep(2)
                log.info("▶️ Ingestion ripresa")
                _set_progress(rc, True, done, total, f"processing-{model_type}")
            
            rel_id = os.path.relpath(path, KB_ROOT)
            filename = os.path.basename(path)
            
            _set_current_doc(rc, {
                "filename": filename,
                "path": path,
                "step": "reading"
            })
            
            try:
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "extracting"})
                text = _read_text(path)
                
                if not text:
                    log.debug(f"⚠️ Nessun testo estratto da {filename}")
                    done += 1
                    continue
                
                title = filename
                metadata = extract_metadata(path, KB_ROOT)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "postgres"})
                cur.execute("""
                    INSERT INTO documents (
                        id, path, title, content,
                        ext, area, anno, cliente, oggetto, tipo_doc,
                        codice_appalto, categoria, versione
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET
                        path=EXCLUDED.path,
                        title=EXCLUDED.title,
                        content=EXCLUDED.content,
                        ext=EXCLUDED.ext,
                        area=EXCLUDED.area,
                        anno=EXCLUDED.anno,
                        cliente=EXCLUDED.cliente,
                        oggetto=EXCLUDED.oggetto,
                        tipo_doc=EXCLUDED.tipo_doc,
                        codice_appalto=EXCLUDED.codice_appalto,
                        categoria=EXCLUDED.categoria,
                        versione=EXCLUDED.versione,
                        mtime=NOW()
                """, (
                    rel_id, path, title, text,
                    metadata.get('ext'),
                    metadata.get('area'),
                    metadata.get('anno'),
                    metadata.get('cliente'),
                    metadata.get('oggetto'),
                    metadata.get('tipo_doc'),
                    metadata.get('codice_appalto'),
                    metadata.get('categoria'),
                    metadata.get('versione')
                ))
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "chunking"})
                chunks = _chunk_text(text, chunk_size=1500, overlap=200)
                
                if not chunks:
                    done += 1
                    continue
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "embedding", "details": f"{len(chunks)} chunks ({DEVICE})"})
                embeddings = embedder(chunks)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "qdrant"})
                points = []
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    point_id = _generate_point_id(rel_id, i)
                    
                    points.append(PointStruct(
                        id=point_id,
                        vector=embedding,
                        payload={
                            "doc_id": rel_id,
                            "chunk_id": i,
                            "text": chunk,
                            "title": title,
                            "path": path,
                            **{k: v for k, v in metadata.items() if v is not None}
                        }
                    ))
                
                qd.upsert(collection_name=collection_name, points=points)
                
                _set_current_doc(rc, {"filename": filename, "path": path, "step": "meili"})
                meili_batch.append({
                    "id": rel_id,
                    "path": path,
                    "title": title,
                    "content": text[:5000],
                    **{k: v for k, v in metadata.items() if v is not None}
                })
                
                total_chunks += len(chunks)
                
                _add_processing_log(rc, {
                    "filename": filename,
                    "status": "success",
                    "chunks": len(chunks),
                    "meili_indexed": True,
                    "qdrant_vectorized": True,
                    "device": DEVICE
                })
                
                _update_stats(rc,
                    success=done + 1,
                    chunked=total_chunks,
                    qdrant_vectorized=done + 1
                )
            
            except Exception as e:
                log.error(f"❌ Errore processing {filename}: {e}")
                import traceback
                log.error(traceback.format_exc())
                
                _push_failed(rc, {"path": path, "filename": filename, "error": str(e)})
                _add_processing_log(rc, {
                    "filename": filename,
                    "status": "error",
                    "error": str(e)
                })
                _update_stats(rc, failed=rc.llen(Q_REDIS_KEY_FAILED))
            
            finally:
                done += 1
                
                if len(meili_batch) >= 50:
                    try:
                        idx.add_documents(meili_batch)
                        _update_stats(rc, meili_indexed=done)
                    except Exception as e:
                        log.error(f"Errore batch Meilisearch: {e}")
                    meili_batch = []
                
                _set_progress(rc, True, done, total, f"processing-{model_type}")
        
        if meili_batch:
            try:
                idx.add_documents(meili_batch)
                _update_stats(rc, meili_indexed=done)
            except Exception as e:
                log.error(f"Errore final batch: {e}")
    
    _set_current_doc(rc, None)
    _set_progress(rc, False, total, total, "done")
    
    log.info(f"✅ Ingestion completata: {total} documenti, {total_chunks} chunks, modello {model_type}, device {DEVICE}")
    
    # Clear GPU cache finale
    if GPU_AVAILABLE:
        clear_gpu_cache()
    
    return {
        "ok": True,
        "mode": mode,
        "model": model_type,
        "collection": collection_name,
        "total": total,
        "chunks": total_chunks,
        "device": DEVICE
    }
