#!/usr/bin/env python3
"""
fix_docx_support.py - Patch per aggiungere supporto DOCX e altri formati Office

Questo script aggiorna worker/worker_tasks.py per supportare:
- DOCX (tramite python-docx)
- DOC, XLS, XLSX, PPT, PPTX (tramite LibreOffice)

Uso:
    python3 fix_docx_support.py

Output:
    worker/worker_tasks_NEW.py (da rinominare manualmente dopo test)
"""

import os
import sys

NEW_EXTRACT_TEXT = '''def _read_text(path: str) -> str:
    """
    Estrae testo da vari formati di documento.
    
    Supporto:
    - Testo: .txt, .md, .csv, .log
    - PDF: .pdf (via pdftotext)
    - Word: .docx (via python-docx), .doc (via LibreOffice)
    - Excel: .xls, .xlsx (via LibreOffice)
    - PowerPoint: .ppt, .pptx (via LibreOffice)
    
    Returns:
        str: Testo estratto dal documento
    
    Raises:
        RuntimeError: Se estrazione fallisce o formato non supportato
    """
    ext = os.path.splitext(path)[1].lower()
    
    # Formati testo puro
    if ext in (".txt", ".md", ".csv", ".log"):
        with open(path, "r", errors="ignore") as f:
            return f.read()
    
    # PDF via pdftotext
    if ext == ".pdf":
        return _pdftotext(path)
    
    # DOCX via python-docx (veloce e preciso)
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            # Estrai testo da paragrafi e tabelle
            paragraphs = [para.text for para in doc.paragraphs]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    tables_text.append("\\t".join(cell.text for cell in row.cells))
            all_text = "\\n".join(paragraphs + tables_text)
            return all_text
        except ImportError:
            # Fallback a LibreOffice se python-docx non disponibile
            return _libreoffice_convert(path)
        except Exception as e:
            raise RuntimeError(f"Errore estrazione DOCX: {e}")
    
    # Formati Office legacy/altri via LibreOffice
    if ext in (".doc", ".xls", ".xlsx", ".ppt", ".pptx", ".odt", ".ods", ".odp"):
        return _libreoffice_convert(path)
    
    # Formato non supportato
    raise RuntimeError(f"Formato non supportato: {ext}")


def _libreoffice_convert(path: str) -> str:
    """
    Converte documenti Office in testo usando LibreOffice headless.
    
    Args:
        path: Percorso file da convertire
    
    Returns:
        str: Testo estratto
    
    Raises:
        RuntimeError: Se conversione fallisce
    """
    import tempfile
    import subprocess
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Converti in TXT
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "txt:Text",
            "--outdir", tmpdir,
            path
        ], capture_output=True, text=True, timeout=60)
        
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversione fallita: {result.stderr}")
        
        # Leggi file TXT generato
        basename = os.path.splitext(os.path.basename(path))[0]
        txt_path = os.path.join(tmpdir, f"{basename}.txt")
        
        if not os.path.exists(txt_path):
            raise RuntimeError(f"File convertito non trovato: {txt_path}")
        
        with open(txt_path, "r", errors="ignore") as f:
            return f.read()
'''

def patch_worker_file():
    """Applica patch al file worker/worker_tasks.py"""
    
    worker_file = "worker/worker_tasks.py"
    
    if not os.path.exists(worker_file):
        print(f"❌ File non trovato: {worker_file}")
        print("   Assicurati di eseguire questo script dalla root del repository")
        return False
    
    print(f"📖 Lettura {worker_file}...")
    with open(worker_file, "r") as f:
        content = f.read()
    
    # Trova la funzione _read_text esistente
    import re
    pattern = r'def _read_text\(path: str\) -> str:.*?(?=\ndef |\nclass |\Z)'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        print("❌ Funzione _read_text non trovata nel file")
        return False
    
    old_function = match.group(0)
    print(f"✅ Trovata funzione _read_text (lunghezza: {len(old_function)} char)")
    
    # Sostituisci con nuova versione
    new_content = content.replace(old_function, NEW_EXTRACT_TEXT)
    
    # Salva come nuovo file
    output_file = "worker/worker_tasks_NEW.py"
    with open(output_file, "w") as f:
        f.write(new_content)
    
    print(f"✅ Patch applicata! File salvato in: {output_file}")
    print()
    print("🔍 Verifica le modifiche:")
    print(f"   diff worker/worker_tasks.py {output_file}")
    print()
    print("✅ Se tutto ok, rinomina il file:")
    print(f"   mv {output_file} worker/worker_tasks.py")
    print()
    print("⚠️  Nota: Richiede dipendenze:")
    print("   - python-docx (già in requirements)")
    print("   - libreoffice (già in Dockerfile_worker)")
    
    return True


def create_test_script():
    """Crea script di test per verificare supporto formati"""
    
    test_script = '''#!/usr/bin/env python3
"""
test_extract_text.py - Test estrazione testo da vari formati

Uso:
    python3 test_extract_text.py /path/to/test/files/
"""

import sys
import os
from pathlib import Path

# Importa funzione modificata
sys.path.insert(0, 'worker')
from worker_tasks import _read_text

def test_extraction(test_dir):
    """Testa estrazione da tutti i file nella directory"""
    
    test_dir = Path(test_dir)
    if not test_dir.exists():
        print(f"❌ Directory non trovata: {test_dir}")
        return
    
    # Trova file di test
    files = [f for f in test_dir.rglob("*") if f.is_file()]
    
    if not files:
        print(f"⚠️  Nessun file trovato in {test_dir}")
        return
    
    print(f"🧪 Test estrazione su {len(files)} file\\n")
    
    results = {
        "success": [],
        "failed": []
    }
    
    for filepath in files:
        ext = filepath.suffix.lower()
        try:
            text = _read_text(str(filepath))
            text_len = len(text)
            preview = text[:100].replace("\\n", " ")
            
            print(f"✅ {filepath.name}")
            print(f"   Ext: {ext}, Lunghezza: {text_len} char")
            print(f"   Preview: {preview}...")
            print()
            
            results["success"].append((filepath.name, ext, text_len))
            
        except Exception as e:
            print(f"❌ {filepath.name}")
            print(f"   Errore: {e}")
            print()
            results["failed"].append((filepath.name, ext, str(e)))
    
    # Riepilogo
    print("\\n" + "="*60)
    print("📊 RIEPILOGO TEST")
    print("="*60)
    print(f"✅ Successi: {len(results['success'])}")
    print(f"❌ Fallimenti: {len(results['failed'])}")
    
    if results["success"]:
        print("\\n✅ File estratti con successo:")
        for name, ext, length in results["success"]:
            print(f"   - {name} ({ext}): {length} char")
    
    if results["failed"]:
        print("\\n❌ File falliti:")
        for name, ext, error in results["failed"]:
            print(f"   - {name} ({ext}): {error}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python3 test_extract_text.py /path/to/test/files/")
        sys.exit(1)
    
    test_extraction(sys.argv[1])
'''
    
    test_file = "test_extract_text.py"
    with open(test_file, "w") as f:
        f.write(test_script)
    
    os.chmod(test_file, 0o755)
    print(f"✅ Script di test creato: {test_file}")
    print(f"   Uso: python3 {test_file} /path/to/test/files/")


if __name__ == "__main__":
    print("🔧 Fix Supporto DOCX e Formati Office")
    print("="*50)
    print()
    
    success = patch_worker_file()
    
    if success:
        print()
        create_test_script()
        print()
        print("✨ Completato!")
    else:
        print()
        print("❌ Patch fallita. Verifica manualmente il file.")
        sys.exit(1)
